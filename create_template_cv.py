import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
import os
import shutil
from docx2pdf import convert

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=160, right=160):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def remove_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/><w:insideH w:val="none"/><w:insideV w:val="none"/></w:tblBorders>')
    tblPr.append(borders)

doc = Document()

# Page setup - optimized to fit cleanly on 1 page
for s in doc.sections:
    s.top_margin = Inches(0.28)
    s.bottom_margin = Inches(0.25)
    s.left_margin = Inches(0.28)
    s.right_margin = Inches(0.28)
    s.page_width = Inches(8.5)
    s.page_height = Inches(11.0)

NAVY = RGBColor(16, 44, 87)       # Deep Navy #102C57
TITLE_NAVY = RGBColor(24, 60, 110)
BODY_GRAY = RGBColor(45, 50, 60)
MUTED_GRAY = RGBColor(95, 100, 110)
ACCENT_BLUE = RGBColor(25, 85, 155)

table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
remove_borders(table)

row = table.rows[0]
cell_left = row.cells[0]
cell_right = row.cells[1]

cell_left.width = Inches(2.45)
cell_right.width = Inches(5.45)

set_cell_background(cell_left, 'F0F4F8')  # Sidebar background tint
set_cell_margins(cell_left, top=120, bottom=120, left=150, right=150)
set_cell_margins(cell_right, top=120, bottom=120, left=200, right=100)

# --- LEFT SIDEBAR CONTENT ---
p_img = cell_left.paragraphs[0]
p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_img.paragraph_format.space_after = Pt(8)
p_img.add_run().add_picture(r'd:\project\maneesh_glass_photo.png', width=Inches(1.65))

def add_sidebar_heading(text):
    p = cell_left.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(1.5)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(8.8)
    r.font.color.rgb = NAVY
    return p

def add_sidebar_item(label, value=None):
    if label:
        p_lbl = cell_left.add_paragraph()
        p_lbl.paragraph_format.space_before = Pt(4.5)
        p_lbl.paragraph_format.space_after = Pt(0.5)
        r_l = p_lbl.add_run(label)
        r_l.bold = True
        r_l.font.size = Pt(8.2)
        r_l.font.color.rgb = TITLE_NAVY
    if value:
        p_val = cell_left.add_paragraph()
        p_val.paragraph_format.space_before = Pt(0)
        p_val.paragraph_format.space_after = Pt(2)
        r_v = p_val.add_run(value)
        r_v.font.size = Pt(7.8)
        r_v.font.color.rgb = BODY_GRAY

add_sidebar_item('Address:', 'Flat no 404, Shree Azad SRA CHS LTD, Azad Link Road, Sanjay Nagar, Malad East, Mumbai')
add_sidebar_item('Date of Birth:', 'December 20, 1985')
add_sidebar_item('Contact:', '+91 9819253815')
add_sidebar_item('Email:', 'maneeshp1@gmail.com\npandeymaneesh@ymail.com')
add_sidebar_item('Personal Details:', 'Married  |  Male  |  English, Hindi')

add_sidebar_heading('Qualification:')
add_sidebar_item('U.P. Board – Graduation (B.A.)', 'March 2007 • 2nd Class')
add_sidebar_item('U.P. Board – H.S.C (12th)', 'March 2004 • 2nd Class')
add_sidebar_item('U.P. Board – S.S.C (10th)', 'March 2002 • 2nd Class')

add_sidebar_heading('Certifications (2026):')
add_sidebar_item('NISM Series 25A', 'Certified in 2026')
add_sidebar_item('IRDA Examination', 'Certified in 2026 (Life & General)')

add_sidebar_heading('Core Competencies:')
skills = [
    'Digital Broking & Demat Acquisition',
    'Team Leadership (50 RMs & 5 TLs)',
    'Multi-Product Revenue & MIP Strategy',
    'Equity, Derivatives, MF & Wealth',
    'Recruitment, Training & Mentoring',
    'MIS Reporting, CRM & Dialler Software'
]
for sk in skills:
    p_sk = cell_left.add_paragraph(style='List Bullet')
    p_sk.paragraph_format.space_before = Pt(0)
    p_sk.paragraph_format.space_after = Pt(1)
    p_sk.paragraph_format.left_indent = Inches(0.15)
    r_sk = p_sk.add_run(sk)
    r_sk.font.size = Pt(7.5)
    r_sk.font.color.rgb = BODY_GRAY

# --- RIGHT MAIN CONTENT ---
p_name = cell_right.paragraphs[0]
p_name.paragraph_format.space_before = Pt(0)
p_name.paragraph_format.space_after = Pt(1)
r_name = p_name.add_run('MANEESH KUMAR PANDEY')
r_name.bold = True
r_name.font.size = Pt(18)
r_name.font.color.rgb = NAVY

p_sub = cell_right.add_paragraph()
p_sub.paragraph_format.space_after = Pt(6)
r_sub = p_sub.add_run('Regional Business Head (Digital Business) | Broking & Financial Services')
r_sub.bold = True
r_sub.font.size = Pt(9.2)
r_sub.font.color.rgb = ACCENT_BLUE

def add_main_heading(title):
    p = cell_right.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(9.8)
    r.font.color.rgb = NAVY
    return p

# Profile Summary
add_main_heading('Profile Summary')
p_sum = cell_right.add_paragraph('Results-driven Financial Services Leader with 17+ years of leadership experience across Digital Broking, Wealth Advisory, Demat acquisition, and Multi-product revenue generation. Adept at managing large sales teams (50 RMs & 5 TLs), formulating high-impact MIP strategies, talent grooming, and consistently exceeding business targets.')
p_sum.paragraph_format.space_after = Pt(4)
for r in p_sum.runs:
    r.font.size = Pt(8.1)
    r.font.color.rgb = BODY_GRAY

# Key Achievements
add_main_heading('Key Achievements:')
achievements = [
    'Pan India 2nd Best Performer for outstanding sales leadership and national revenue generation.',
    'Appreciation Letter from Executive Director of Angel Broking Ltd for business performance.',
    'Qualified Pan India contest for Wealth Management and won international foreign trip reward (2014).',
    'Leading 50 RMs & 5 Team Leaders at HDFC Securities Ltd with consistent month-on-month target achievement.'
]
for ach in achievements:
    p_ach = cell_right.add_paragraph(style='List Bullet')
    p_ach.paragraph_format.space_before = Pt(0)
    p_ach.paragraph_format.space_after = Pt(1.2)
    p_ach.paragraph_format.left_indent = Inches(0.18)
    r_ach = p_ach.add_run(ach)
    r_ach.font.size = Pt(8.0)
    r_ach.font.color.rgb = BODY_GRAY

# Professional Experience
add_main_heading('Professional Experience:')

def add_main_job(company_role, duration, bullets):
    p_jr = cell_right.add_paragraph()
    p_jr.paragraph_format.space_before = Pt(5)
    p_jr.paragraph_format.space_after = Pt(1)
    
    r_head = p_jr.add_run(f'{company_role} | {duration}')
    r_head.bold = True
    r_head.font.size = Pt(8.6)
    r_head.font.color.rgb = NAVY
    
    for b in bullets:
        p_b = cell_right.add_paragraph(style='List Bullet')
        p_b.paragraph_format.space_before = Pt(0)
        p_b.paragraph_format.space_after = Pt(1)
        p_b.paragraph_format.left_indent = Inches(0.18)
        r_b = p_b.add_run(b)
        r_b.font.size = Pt(7.8)
        r_b.font.color.rgb = BODY_GRAY

add_main_job(
    'HDFC Securities Ltd – Regional Business Head (Digital Business)',
    'Jul 2021 – Present',
    [
        'Leading a high-performing team of 50 Relationship Managers (RMs) and 5 Team Leaders.',
        'Managing online lead-based Demat account acquisition, digital onboarding, and activation.',
        'Driving multi-segment revenue across AOC, Value Plan, Brokerage, Life & Health Insurance, and Paid Baskets.',
        'Recruiting RMs & TLs, providing training on Equity/Wealth, and executing Monthly Incentive Plans (MIP).',
        'Conducting daily KPI reviews and driving the region toward revenue targets and top rankings.'
    ]
)

add_main_job(
    'IIFL Securities Ltd – Area Sales Manager',
    'Aug 2018 – Jun 2021',
    [
        'Supervised 20 sales associates with 2 Sales Managers; spearheaded sales improvement initiatives.',
        'Trained and coached executives/SMs to consistently hit volume, revenue, and client satisfaction targets.'
    ]
)

add_main_job(
    'Kunvarji Finstock Pvt. Ltd – Area Sales Manager',
    'Jan 2018 – Aug 2018',
    [
        'Organized investor awareness activities for Equity, Commodities, Currency, and PMS.',
        'Monitored revenue & margin reports, expanded business channels, and led sales recruitment.'
    ]
)

add_main_job(
    'Angel Broking Ltd – Area Sales Manager / Sales Manager / Unit Manager / Sr. RE',
    'Jan 2008 – Jan 2018',
    [
        'Progressive 10-year promotion track: Sr. Relationship Executive -> Unit Manager -> Sales Manager -> ASM.',
        'Drove broking, PMS, and insurance distribution; maintained high client engagement and performance awards.'
    ]
)

add_main_job(
    'India Infoline Ltd – Relationship Executive',
    'Jul 2007 – Dec 2007',
    [
        'Lead generation, Demat/Trading account opening, and retail customer advisory support.'
    ]
)

output_docx = r'd:\project\Maneesh_Pandey_WA_Template_CV.docx'
output_pdf = r'd:\project\Maneesh_Pandey_WA_Template_CV.pdf'
down_docx = r'C:\Users\pande\Downloads\Maneesh_Pandey_WA_Template_CV.docx'
down_pdf = r'C:\Users\pande\Downloads\Maneesh_Pandey_WA_Template_CV.pdf'

doc.save(output_docx)
doc.save(down_docx)
print('DOCX saved.')

convert(output_docx, output_pdf)
shutil.copy2(output_pdf, down_pdf)
print('PDF generated successfully.')
