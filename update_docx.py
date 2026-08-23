import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def remove_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/><w:insideH w:val="none"/><w:insideV w:val="none"/></w:tblBorders>')
    tblPr.append(borders)

doc = Document()

for s in doc.sections:
    s.top_margin = Inches(0.4)
    s.bottom_margin = Inches(0.4)
    s.left_margin = Inches(0.4)
    s.right_margin = Inches(0.4)
    s.page_width = Inches(8.5)
    s.page_height = Inches(11.0)

NAVY = RGBColor(16, 44, 87)
TITLE_NAVY = RGBColor(24, 60, 110)
BODY_GRAY = RGBColor(50, 55, 65)
ACCENT_BLUE = RGBColor(30, 90, 160)

table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
remove_borders(table)

row = table.rows[0]
cell_left = row.cells[0]
cell_right = row.cells[1]

cell_left.width = Inches(2.6)
cell_right.width = Inches(5.1)

set_cell_background(cell_left, 'F2F6FA')
set_cell_margins(cell_left, top=180, bottom=180, left=180, right=180)
set_cell_margins(cell_right, top=180, bottom=180, left=240, right=140)

# Sidebar
p_img = cell_left.paragraphs[0]
p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_img.paragraph_format.space_after = Pt(12)
p_img.add_run().add_picture(r'd:\project\maneesh_glass_photo.png', width=Inches(1.8))

def add_sb_head(txt):
    p = cell_left.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = NAVY

def add_sb_item(lbl, val=None):
    if lbl:
        p = cell_left.add_paragraph()
        p.paragraph_format.space_before = Pt(4.5)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(lbl)
        r.bold = True
        r.font.size = Pt(8.8)
        r.font.color.rgb = TITLE_NAVY
    if val:
        p = cell_left.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3.5)
        r = p.add_run(val)
        r.font.size = Pt(8.2)
        r.font.color.rgb = BODY_GRAY

add_sb_head('CONTACT DETAILS')
add_sb_item('Address:', 'Flat No. 404, Shree Azad SRA CHS Ltd, Azad Link Road, Sanjay Nagar Pathanwadi, Rani Sati Marg, Malad East, Mumbai – 400097')
add_sb_item('Mobile / WhatsApp:', '+91 9819253815')
add_sb_item('Email Addresses:', 'maneeshp1@gmail.com\npandeymaneesh@ymail.com')
add_sb_item('Date of Birth:', 'December 20, 1985')
add_sb_item('Personal Information:', 'Gender: Male\nMarital Status: Married\nFather: Indramani Pandey\nLanguages: English, Hindi')

add_sb_head('SKILLS & CORE CAPABILITIES')
skills = [
    'Leadership & Team Development',
    'Client Servicing & Relationship Management',
    'Business Improvement & Strategic Execution',
    'Audit, SLA Compliance & Metrics',
    'Critical Thinking & Talent Recruitment',
    'Digital Investment Platforms & Demat Growth',
    'Multi-Product Cross-Sell (MF, GI, LI, PMS)',
    'RM Productivity Enhancement & MIP Plans'
]
for s in skills:
    p = cell_left.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.left_indent = Inches(0.18)
    r = p.add_run(s)
    r.font.size = Pt(8.0)
    r.font.color.rgb = BODY_GRAY

add_sb_head('QUALIFICATION')
add_sb_item('Graduation (B.A.):', 'U.P. Board | March 2007 • Second Class')
add_sb_item('H.S.C. (12th):', 'U.P. Board | March 2004 • Second Class')
add_sb_item('S.S.C. (10th):', 'U.P. Board | March 2002 • Second Class')

add_sb_head('CERTIFICATIONS (2026)')
add_sb_item('NISM Series 25A Certification:', '• Completed in Year 2026\nComprehensive securities market advisory, operations & compliance certification.')
add_sb_item('IRDA Examination:', '• Completed in Year 2026\nInsurance Regulatory and Development Authority Certified for Life & General Insurance.')

add_sb_head('TECHNICAL & TOOL EXPERTISE')
tech = [
    'CRM & LMS Platforms: Lead tracking, funnels',
    'Dialling Systems: Auto-diallers, predictive telephony',
    'Trading Platforms: ODIN Diet, Omnesys, Mobile Apps',
    'Office Suite: Advanced MS Excel, Word, PowerPoint',
    'Operating Systems: Windows XP, Win 7, Win 10, Win 11'
]
for t in tech:
    p = cell_left.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.left_indent = Inches(0.18)
    r = p.add_run(t)
    r.font.size = Pt(8.0)
    r.font.color.rgb = BODY_GRAY

add_sb_head('LEADERSHIP ATTRIBUTES')
lead = ['High-Energy Team Leadership', 'Strategic Sales Forecasting', 'Data-Driven MIS Analytics', 'Dispute & Grievance Redressal', 'Self-Initiative & Persistence', 'Seeking Continuous Improvement', 'Client Retention & Trust Building']
for l in lead:
    p = cell_left.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.left_indent = Inches(0.18)
    r = p.add_run(l)
    r.font.size = Pt(8.0)
    r.font.color.rgb = BODY_GRAY

add_sb_head('HOBBIES & INTERESTS')
hob = ['Interacting With People & Networking', 'Reading Newspapers & Financial Press', 'Listening to Music', 'Continuous Skill Learning']
for h in hob:
    p = cell_left.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.left_indent = Inches(0.18)
    r = p.add_run(h)
    r.font.size = Pt(8.0)
    r.font.color.rgb = BODY_GRAY

# Main Column
p_name = cell_right.paragraphs[0]
p_name.paragraph_format.space_before = Pt(0)
p_name.paragraph_format.space_after = Pt(2)
r_name = p_name.add_run('MANEESH KUMAR PANDEY')
r_name.bold = True
r_name.font.size = Pt(20)
r_name.font.color.rgb = NAVY

p_sub = cell_right.add_paragraph()
p_sub.paragraph_format.space_after = Pt(10)
r_sub = p_sub.add_run('Regional Business Head (Digital Business) | Broking & Financial Services')
r_sub.bold = True
r_sub.font.size = Pt(10)
r_sub.font.color.rgb = ACCENT_BLUE

def add_mn_head(txt):
    p = cell_right.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = NAVY

add_mn_head('PROFILE SUMMARY')
p_sum = cell_right.add_paragraph("Seasoned financial services professional with 19+ years' leadership experience across Broking, Digital Wealth Advisory, and Multi-Product Distribution, specializing in sales capability development, digital investment platforms, RM productivity enhancement, quality audits, team recruitment, and large sales force leadership (50+ Relationship Managers & 5 Team Leaders). Proven track record in scaling digital Demat acquisition, executing high-yield Monthly Incentive Plans (MIP), driving multi-segment revenues across Equity, Derivatives, Mutual Funds, and Insurance, and consistently exceeding business targets.")
p_sum.paragraph_format.space_after = Pt(6)
for r in p_sum.runs:
    r.font.size = Pt(8.8)
    r.font.color.rgb = BODY_GRAY

add_mn_head('KEY ACHIEVEMENTS')
ach = [
    'Pan India 2nd Best Performer Award: Conferred national award for superlative business volume, team activation, and revenue generation across all regional units.',
    "Executive Director's Letter of Appreciation: Received formal commendation letter from Executive Director, Angel Broking Ltd for extraordinary sales leadership and consistent target exceedance.",
    'Pan India Wealth Management Contest Winner (2014): Qualified nationwide Wealth Management contest and achieved sponsored foreign trip reward.',
    'Trained & Mentored 500+ Sales Professionals: Successfully developed sales capabilities, digital platform proficiency, and product knowledge across banking and broking channels.',
    'Led High-Capacity Business Units: Managed sales units with 150+ members historically and currently commanding 50 RMs + 5 TLs digital acquisition unit at HDFC Securities.'
]
for a in ach:
    p = cell_right.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run(a)
    r.font.size = Pt(8.5)
    r.font.color.rgb = BODY_GRAY

add_mn_head('PROFESSIONAL EXPERIENCE')

def add_exp(comp_role, dur, bullets):
    p = cell_right.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(1.5)
    r1 = p.add_run(f'{comp_role} | {dur}')
    r1.bold = True
    r1.font.size = Pt(9.2)
    r1.font.color.rgb = NAVY
    for b in bullets:
        pb = cell_right.add_paragraph(style='List Bullet')
        pb.paragraph_format.space_before = Pt(0)
        pb.paragraph_format.space_after = Pt(2)
        pb.paragraph_format.left_indent = Inches(0.2)
        rb = pb.add_run(b)
        rb.font.size = Pt(8.5)
        rb.font.color.rgb = BODY_GRAY

add_exp(
    'HDFC Securities Ltd — Regional Business Head (Digital Business)',
    'July 2021 – Present | Mumbai',
    [
        'Team Leadership & Span of Control: Managing a team size of 50 Relationship Managers (RMs) supervised through 5 Team Leaders (TLs) in Demat Account Business Generation and digital acquisition.',
        'Digital Demat Acquisition: Managing online lead-based Demat account acquisition and ensuring consistent account opening, onboarding, and activation.',
        'Multi-Segment Revenue Driving: Driving team revenue through multiple segments, including Account Opening Charges (AOC), Value Plan, Brokerage (Equity & Derivatives), Mutual Funds, Life Insurance, Health Insurance, and Paid Stock Baskets.',
        'Target Planning & Execution: Responsible for achieving daily, monthly, and overall revenue targets through effective team planning and strategic execution.',
        'Sales Force Recruitment & Training: Recruiting new RMs & TLs and providing comprehensive need-based training on Equity, Derivatives, Wealth, and specialized financial products.',
        'MIP Strategies & Contests: Creating and implementing MIP (Monthly Incentive Plan) strategies and performance contests for the team to motivate RMs, increase productivity, and boost business.',
        'Performance Reviews & MIS Analysis: Conducting daily performance reviews to track individual and team-level KPIs, analyzing MIS reports, and acting accordingly on bottom performer RMs & TLs to improve output.',
        'Strategic Tech & Marketing Collaboration: Actively involved in key meetings like LMS (Lead Management System), Dialling software, Marketing team for lead generation, and CRM.',
        'Client Feedback Desk: Daily calling to new clients for taking their feedback and suggestions; compiling and analyzing sales figures.'
    ]
)

add_exp(
    'IIFL Securities Ltd — Area Sales Manager',
    '20 August 2018 – 2 June 2021 | Mumbai',
    [
        'Sales Force Supervision: Supervised a sales force of 20 sales associates with 2 Sales Managers (SMs).',
        'Cross-Functional Initiatives: Spearheaded cross-functional initiatives to achieve sales improvement, revenue growth, and category strategy changes.',
        'Coaching & Mentoring: Trained, coached, and mentored executives and Sales Managers to consistently achieve volume, revenue, and product targets.',
        'Product & Process Training: Delivered product and process training, client need analysis, and digital platform demos across Equity, Mutual Funds, Bonds & Derivatives.',
        'Stakeholder & Customer Relationship Management: Customer-oriented relationship management; managed customer relationships and business stakeholders, executed joint collaborative initiatives, and monitored sales team.',
        'SOPs & Policy Implementation: Facilitated the development, implementation, and maintenance of processes, policies, guidelines, Standard Operating Procedures (SOPs), and Business Operating Principles.'
    ]
)

add_exp(
    'Kunvarji Finstock Pvt. Ltd — Area Sales Manager',
    '9 January 2018 – 8 August 2018 | Mumbai',
    [
        'Investor Awareness Activities: Organized multiple awareness activities like investor meets for Equity, Commodities, and Currency, plus promotional activities in corporate environments.',
        'Product & Advisory Awareness: Created awareness regarding Equity, Currency, Commodity, and Portfolio Management Services (PMS), and also selling insurance products.',
        'Margin & Revenue Reports: Maintained Margin reports and Revenue reports of the team; created multiple reports to represent as a team in front of senior management.',
        'Team Grooming & Training: Conducted training for sales team to groom them for sales activities; set goals for individuals as well as the team.',
        'Omni-Channel Lead Pipelines: Increased business opportunities through various routes to market like SMS campaigns, email campaigns, referral programs for clients, etc.',
        'Recruitment & Territory Expansion: Handled team recruitment and expansion of business; established, maintained, and expanded customer base for the organization.',
        'Sales Strategy & Contests: Developed sales strategies to meet targets; designed different contests for the sales team and monitored performance to reach goals.',
        'Customer Portfolio Diversification: Serviced the needs of existing customers and helped them diversify their investment portfolios.'
    ]
)

add_exp(
    'Angel Broking Ltd — Area Sales Manager (Total 10 Years: 2008 – 2018)',
    '1 October 2017 – 8 January 2018 | Mumbai',
    [
        'Awareness & Multi-Product Sales: Creating awareness regarding Equity, Currency, Commodity, and Portfolio Management Services, and also selling insurance products.',
        'Target Achievement & Escalation Desk: Reaching the targets and goals set for the area; managing & resolving escalation queries of customers.',
        'Client Feedback & Engagement: Daily calling to new clients for taking their feedback and suggestions; executing joint collaborative initiatives and monitoring sales team.'
    ]
)

add_exp(
    'Angel Broking Ltd — Sales Manager (Promoted from Unit Manager)',
    'January 2014 – September 2017 | Mumbai',
    [
        'Team Sales Delivery: Creating awareness regarding Equity, Currency, Commodity, and Portfolio Management Services (PMS) and also selling insurance products from team.',
        'Target Setting & Achievement: Setting up targets for the team and motivating them to achieve monthly milestones; maintaining healthy relations between clients and team members.',
        'Promotion Milestone: Got promoted from Unit Manager to Sales Manager post based on consistent high performance.'
    ]
)

add_exp(
    'Angel Broking Ltd — Unit Manager (Promoted from Sr. RE)',
    'June 2010 – December 2013 | Mumbai',
    [
        'Business Unit Expansion: Creating awareness regarding Equity, Currency, Commodity, and PMS, and selling insurance products from the team; setting targets and achieving quota.',
        'Trading Demos & Software Training: Trained sales representatives on online trading platforms (ODIN/Diet Odin) and product/process systems.',
        'Promotion Milestone: Got promoted from Senior Relationship Executive to Unit Manager post.'
    ]
)

add_exp(
    'Angel Broking Ltd — Sr. Relationship Executive',
    '7 January 2008 – December 2010 | Mumbai',
    [
        'Cross-Selling & Demat Acquisition: Cross-selling insurance products along with opening of Demat accounts; achieving the set targets consistently.',
        'Demos & Query Coordination: Providing demos to clients for trading products; resolving issues by coordinating with the respective dealer & advisory of the client.'
    ]
)

add_exp(
    'India Infoline Ltd (IIFL) — Relationship Executive',
    '14 July 2007 – December 2007 | Mumbai',
    [
        'Lead Generation & Account Opening: Generating new leads; opening of Trading & Demat accounts for retail investors.',
        'Client Support & Conversion: Providing after-sales service to existing clients; solving issues of prospective clients & converting into active accounts.'
    ]
)

add_mn_head('KEY DOMAIN EXPERTISE & STRATEGIC HIGHLIGHTS')
p_strat = cell_right.add_paragraph('• Digital Tele-Sales & Lead Generation: In-depth expertise in managing LMS (Lead Management Systems), auto-dialling telephony software, CRM, and digital marketing team alignments for maximum acquisition ROI.\n• Regulatory & Compliance Governance: 100% compliant adherence to SEBI, NSE, BSE, NSDL, CDSL, and IRDA guidelines across client KYC, margin collection, and risk disclosure standards.\n• Multi-Product Cross-Sell Architecture: Deep domain expertise in structuring multi-product financial solutions encompassing Equity, F&O, Mutual Funds, Life Insurance, Health Insurance, and PMS.')
for r in p_strat.runs:
    r.font.size = Pt(8.5)
    r.font.color.rgb = BODY_GRAY

p_sig = cell_right.add_paragraph()
p_sig.paragraph_format.space_before = Pt(8)
p_sig.add_run('Place: Mumbai\nDate: ____________________\n\n').bold = True
r_name_sig = p_sig.add_run('MANEESH KUMAR PANDEY')
r_name_sig.bold = True
r_name_sig.font.size = Pt(9.5)
r_name_sig.font.color.rgb = NAVY

docx_proj = r'd:\project\Maneesh_Pandey_WA_Template_CV.docx'
docx_down = r'C:\Users\pande\Downloads\Maneesh_Pandey_WA_Template_CV.docx'
doc.save(docx_proj)
doc.save(docx_down)
print('DOCX successfully updated with 19+ years experience.')
